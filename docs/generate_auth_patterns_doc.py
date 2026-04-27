"""
Generate docs/Workload-Identity-Auth-Patterns.docx — example .NET projects
illustrating two ways a DKP workload can authenticate to Confluent Cloud
through Microsoft Entra ID:

  Pattern A — Client ID + Client Secret (legacy / on-prem / local dev)
  Pattern B — Azure Function App + Managed Identity (passwordless)

Usage:
    pip install python-docx --break-system-packages
    python docs/generate_auth_patterns_doc.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

OUT = Path(__file__).parent / "Workload-Identity-Auth-Patterns.docx"

CONFLUENT_BLUE = RGBColor(0x00, 0x3E, 0x7E)
CODE_BG        = "F2F2F2"
CALLOUT_BG     = "EAF2FB"
WARN_BG        = "FFF4E5"
MUTED          = RGBColor(0x55, 0x55, 0x55)
WARN           = RGBColor(0xB9, 0x60, 0x00)


# ---------- low-level helpers (same style as the runbook generator) ----------

def shade(element, fill: str) -> None:
    tc_pr = element._tc.get_or_add_tcPr() if hasattr(element, "_tc") \
            else element._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CONFLUENT_BLUE


def add_code_block(doc: Document, code: str, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    shade(cell, CODE_BG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.paragraphs[0].clear()
    for line in code.splitlines() or [""]:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Menlo"
        run.font.size = Pt(9)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]],
                 col_widths: tuple[int, int] = (6, 11)) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text = "Field", "Value"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for i, (k, v) in enumerate(rows, start=1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    tbl.columns[0].width = Cm(col_widths[0])
    tbl.columns[1].width = Cm(col_widths[1])
    doc.add_paragraph()


def add_compare_table(doc: Document, headers: list[str],
                      rows: list[list[str]],
                      col_widths: list[int]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            tbl.rows[ri].cells[ci].text = val
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = Cm(w)
    doc.add_paragraph()


def add_callout(doc: Document, label: str, text: str,
                color: RGBColor = CONFLUENT_BLUE,
                bg: str = CALLOUT_BG) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, bg)
    p = cell.paragraphs[0]
    run_label = p.add_run(f"{label}  ")
    run_label.font.bold = True
    run_label.font.color.rgb = color
    run_text = p.add_run(text)
    run_text.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


# ---------- code snippets (kept here so the doc and the script stay in sync) ----------

PATTERN_A_CSPROJ = """\
<Project Sdk="Microsoft.NET.Sdk">
  <PropertyGroup>
    <OutputType>Exe</OutputType>
    <TargetFramework>net8.0</TargetFramework>
    <Nullable>enable</Nullable>
    <ImplicitUsings>enable</ImplicitUsings>
    <RuntimeIdentifiers>win-x64;linux-x64;osx-arm64;osx-x64</RuntimeIdentifiers>
  </PropertyGroup>
  <ItemGroup>
    <PackageReference Include="Confluent.Kafka"                              Version="2.6.0" />
    <PackageReference Include="Microsoft.Identity.Client"                    Version="4.66.2" />
    <PackageReference Include="Microsoft.Extensions.Configuration"           Version="8.0.0" />
    <PackageReference Include="Microsoft.Extensions.Configuration.Json"      Version="8.0.0" />
    <PackageReference Include="Microsoft.Extensions.Configuration.EnvironmentVariables" Version="8.0.0" />
  </ItemGroup>
  <ItemGroup>
    <None Update="appsettings.json">
      <CopyToOutputDirectory>PreserveNewest</CopyToOutputDirectory>
    </None>
  </ItemGroup>
</Project>
"""

PATTERN_A_APPSETTINGS = """\
{
  "Kafka": {
    "BootstrapServers": "lkc-1o1jkv.eastus.azure.private.confluent.cloud:9092",
    "LogicalCluster":   "lkc-1o1jkv",
    "IdentityPoolId":   "pool-WR5Q9",
    "Topic":            "dev.deal.calc.mergerarb.json",
    "GroupId":          "dk-confluent-dev-mergerarb-madam-group"
  },
  "AzureAd": {
    "TenantId": "7bab0bc1-bb61-48d7-b2d4-79825c2ac6b8",
    "ClientId": "0dcd6973-2175-4c93-9c10-bcdba0df9065"
  }
}
"""

PATTERN_A_PROGRAM = """\
using Confluent.Kafka;
using Microsoft.Extensions.Configuration;
using Microsoft.Identity.Client;

var cfg = new ConfigurationBuilder()
    .AddJsonFile("appsettings.json", optional: false)
    .AddEnvironmentVariables()
    .Build();

string tenantId      = cfg["AzureAd:TenantId"]!;
string clientId      = cfg["AzureAd:ClientId"]!;
string clientSecret  = Environment.GetEnvironmentVariable("AZURE_CLIENT_SECRET")
    ?? throw new InvalidOperationException("AZURE_CLIENT_SECRET is not set.");

string bootstrap     = cfg["Kafka:BootstrapServers"]!;
string logicalCluster= cfg["Kafka:LogicalCluster"]!;
string identityPool  = cfg["Kafka:IdentityPoolId"]!;
string topic         = cfg["Kafka:Topic"]!;

// MSAL — confidential client, client_credentials flow.
var msal = ConfidentialClientApplicationBuilder
    .Create(clientId)
    .WithTenantId(tenantId)
    .WithClientSecret(clientSecret)
    .Build();

// /.default scope -> aud == clientId in the resulting v2 token.
string[] scopes = new[] { $"{clientId}/.default" };

void OnTokenRefresh(IClient client, string _)
{
    var result = msal.AcquireTokenForClient(scopes).ExecuteAsync()
                     .GetAwaiter().GetResult();

    // Confluent identity-pool extensions go on the OAUTHBEARER message.
    var extensions = new Dictionary<string, string>
    {
        ["logicalCluster"] = logicalCluster,
        ["identityPoolId"] = identityPool,
    };

    client.OAuthBearerSetToken(
        result.AccessToken,
        result.ExpiresOn.ToUnixTimeMilliseconds(),
        principalName: "kafka-client",
        extensions);
}

var producerCfg = new ProducerConfig
{
    BootstrapServers = bootstrap,
    SecurityProtocol = SecurityProtocol.SaslSsl,
    SaslMechanism    = SaslMechanism.OAuthBearer,
};

using var producer = new ProducerBuilder<string, string>(producerCfg)
    .SetOAuthBearerTokenRefreshHandler(OnTokenRefresh)
    .Build();

var dr = await producer.ProduceAsync(topic, new Message<string, string>
{
    Key   = $"demo-{DateTime.UtcNow:O}",
    Value = "Hello from Pattern A (client_id + client_secret).",
});
Console.WriteLine($"Delivered to {dr.TopicPartitionOffset}");
producer.Flush(TimeSpan.FromSeconds(10));
"""

PATTERN_A_RUN = """\
# 1. Provide the secret out-of-band — never check it in.
export AZURE_CLIENT_SECRET="<paste the client_secret value here>"

# 2. Run.
dotnet run
"""

PATTERN_B_CSPROJ = """\
<Project Sdk="Microsoft.NET.Sdk">
  <PropertyGroup>
    <TargetFramework>net8.0</TargetFramework>
    <AzureFunctionsVersion>v4</AzureFunctionsVersion>
    <OutputType>Exe</OutputType>
    <Nullable>enable</Nullable>
    <ImplicitUsings>enable</ImplicitUsings>
  </PropertyGroup>
  <ItemGroup>
    <PackageReference Include="Microsoft.Azure.Functions.Worker"                  Version="1.22.0" />
    <PackageReference Include="Microsoft.Azure.Functions.Worker.Sdk"              Version="1.17.4" />
    <PackageReference Include="Microsoft.Azure.Functions.Worker.Extensions.Timer" Version="4.3.1"  />
    <PackageReference Include="Confluent.Kafka"                                   Version="2.6.0"  />
    <PackageReference Include="Azure.Identity"                                    Version="1.13.1" />
  </ItemGroup>
  <ItemGroup>
    <None Update="host.json">
      <CopyToOutputDirectory>PreserveNewest</CopyToOutputDirectory>
    </None>
    <None Update="local.settings.json">
      <CopyToOutputDirectory>PreserveNewest</CopyToOutputDirectory>
      <CopyToPublishDirectory>Never</CopyToPublishDirectory>
    </None>
  </ItemGroup>
</Project>
"""

PATTERN_B_HOST = """\
{
  "version": "2.0",
  "logging": {
    "applicationInsights": {
      "samplingSettings": { "isEnabled": true }
    }
  }
}
"""

PATTERN_B_LOCAL_SETTINGS = """\
{
  "IsEncrypted": false,
  "Values": {
    "AzureWebJobsStorage":          "UseDevelopmentStorage=true",
    "FUNCTIONS_WORKER_RUNTIME":     "dotnet-isolated",

    "BootstrapServers":             "lkc-1o1jkv.eastus.azure.private.confluent.cloud:9092",
    "LogicalCluster":               "lkc-1o1jkv",
    "IdentityPoolId":               "pool-WR5Q9",
    "Topic":                        "dev.deal.calc.mergerarb.json",

    "AzureManagedIdentityClientId": "<UAMI-client-id>"
  }
}
"""

PATTERN_B_PROGRAM = """\
using Microsoft.Azure.Functions.Worker;
using Microsoft.Extensions.Hosting;

var host = new HostBuilder()
    .ConfigureFunctionsWorkerDefaults()
    .Build();

host.Run();
"""

PATTERN_B_FUNCTION = """\
using Azure.Core;
using Azure.Identity;
using Confluent.Kafka;
using Microsoft.Azure.Functions.Worker;
using Microsoft.Extensions.Logging;

public class KafkaProducerFunction
{
    private readonly ILogger _log;
    private readonly TokenCredential _cred;
    private readonly string _audienceClientId;

    public KafkaProducerFunction(ILoggerFactory loggerFactory)
    {
        _log = loggerFactory.CreateLogger<KafkaProducerFunction>();

        // The UAMI's client ID is the audience we ask Entra to issue the token for.
        // Same value goes into the Confluent pool's app_client_ids list.
        _audienceClientId = Environment
            .GetEnvironmentVariable("AzureManagedIdentityClientId")
            ?? throw new InvalidOperationException(
                "AzureManagedIdentityClientId is not set.");

        // ManagedIdentityCredential targets the specific UAMI assigned to the
        // Function App. DefaultAzureCredential is fine for local dev (it falls
        // back to az-cli / VS creds) but be explicit in prod.
        _cred = new ManagedIdentityCredential(_audienceClientId);
    }

    [Function("KafkaProducer")]
    public async Task Run([TimerTrigger("0 */5 * * * *")] TimerInfo timer)
    {
        var bootstrap      = Environment.GetEnvironmentVariable("BootstrapServers")!;
        var logicalCluster = Environment.GetEnvironmentVariable("LogicalCluster")!;
        var identityPool   = Environment.GetEnvironmentVariable("IdentityPoolId")!;
        var topic          = Environment.GetEnvironmentVariable("Topic")!;

        var producerCfg = new ProducerConfig
        {
            BootstrapServers = bootstrap,
            SecurityProtocol = SecurityProtocol.SaslSsl,
            SaslMechanism    = SaslMechanism.OAuthBearer,
        };

        void OnTokenRefresh(IClient client, string _)
        {
            // No client_secret. Identity is asserted by the Azure platform
            // itself via IMDS; this call returns a JWT signed by Entra.
            var token = _cred.GetToken(
                new TokenRequestContext(new[] { $"{_audienceClientId}/.default" }),
                CancellationToken.None);

            var extensions = new Dictionary<string, string>
            {
                ["logicalCluster"] = logicalCluster,
                ["identityPoolId"] = identityPool,
            };

            client.OAuthBearerSetToken(
                token.Token,
                token.ExpiresOn.ToUnixTimeMilliseconds(),
                principalName: "function-app",
                extensions);
        }

        using var producer = new ProducerBuilder<string, string>(producerCfg)
            .SetOAuthBearerTokenRefreshHandler(OnTokenRefresh)
            .Build();

        var dr = await producer.ProduceAsync(topic, new Message<string, string>
        {
            Key   = $"fn-{DateTime.UtcNow:O}",
            Value = "Hello from Pattern B (passwordless)."
        });
        producer.Flush(TimeSpan.FromSeconds(10));
        _log.LogInformation("Published to {Tpo}", dr.TopicPartitionOffset);
    }
}
"""

PATTERN_B_DEPLOY = """\
# 1. Build & publish to Azure (Function App must already exist).
func azure functionapp publish <function-app-name>

# 2. Confirm the UAMI is attached:
#    Azure portal -> Function App -> Identity -> User assigned tab -> UAMI listed.

# 3. Confirm the UAMI's client ID is in the Confluent pool's app_client_ids:
#    See terraform/live/<env>/workloads.json. Example:
#      "positionapi": {
#        "app_client_ids": [
#          "<workload-app-client-id>",
#          "<UAMI-client-id>"     <-- add this
#        ],
#        ...
#      }
#    terragrunt apply
"""

PATTERN_B_FIC_CODE = """\
using Azure.Core;
using Azure.Identity;

string tenantId         = "<tenant-id>";
string uamiClientId     = "<UAMI-client-id>";
string workloadClientId = "<workload-app-client-id>";

// MI gets a federation assertion (audience: api://AzureADTokenExchange).
var miCred = new ManagedIdentityCredential(uamiClientId);

// MSAL (via ClientAssertionCredential) presents that assertion to the
// workload app's /oauth2/v2.0/token endpoint, in exchange for a v2 access
// token issued by the workload app itself.
var cred = new ClientAssertionCredential(
    tenantId,
    workloadClientId,
    async ct =>
    {
        var assertion = await miCred.GetTokenAsync(
            new TokenRequestContext(
                new[] { "api://AzureADTokenExchange/.default" }), ct);
        return assertion.Token;
    });

var token = await cred.GetTokenAsync(
    new TokenRequestContext(new[] { $"{workloadClientId}/.default" }),
    CancellationToken.None);

// token.Token is a v2 JWT with aud = workloadClientId, iss = login.microsoftonline.com/<tid>/v2.0
"""


# ---------- doc body ----------

def build() -> Document:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DK Workload Identity")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = CONFLUENT_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(".NET Authentication Patterns")
    r.font.size = Pt(16); r.font.color.rgb = CONFLUENT_BLUE

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Pattern A — Client ID + Client Secret    •    "
                     "Pattern B — Function App + Managed Identity (passwordless)")
    r.font.size = Pt(11); r.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"v1.0 · {date.today():%B %Y} · Prepared by Confluent Professional Services")
    r.font.size = Pt(10); r.font.color.rgb = MUTED

    doc.add_paragraph()
    doc.add_paragraph()

    # ---- Purpose ----
    add_heading(doc, "Purpose", 1)
    doc.add_paragraph(
        "Two reference .NET projects illustrating how a DKP workload obtains a Microsoft Entra "
        "ID access token and uses it to authenticate to Confluent Cloud over SASL/OAUTHBEARER. "
        "Both projects target the same Confluent identity pool — the only thing that changes "
        "is how the workload proves its identity to Entra:"
    )
    add_bullets(doc, [
        "Pattern A — the workload presents a client ID and client secret. "
        "Use this for legacy on-prem servers, container instances outside Azure, "
        "and local development.",
        "Pattern B — the workload runs as an Azure Function App with a "
        "User-Assigned Managed Identity (UAMI) attached. The Azure platform "
        "asserts the identity directly through the Instance Metadata Service "
        "(IMDS); no secret is stored anywhere. This is the recommended pattern "
        "for any workload running on Azure (Function Apps, Container Apps, "
        "App Service, AKS pods, VMs, etc.).",
    ])

    add_callout(doc, "Audience.",
                "DKP application developers and the Azure / Entra administrators who will "
                "configure the workload identities. Existing knowledge of Confluent.Kafka "
                "for .NET is assumed; the snippets below show only the auth-relevant code.")

    # ---- Patterns at a glance ----
    add_heading(doc, "Patterns at a glance", 1)
    add_compare_table(
        doc,
        headers=["Aspect", "Pattern A — Secret", "Pattern B — Managed Identity"],
        rows=[
            ["How identity is proven to Entra",
             "Code presents client_id + client_secret over HTTPS.",
             "Azure platform injects the identity via IMDS (only reachable inside the host)."],
            ["Where the secret lives",
             "Env var, Key Vault, or App Setting.",
             "Nowhere. There is no shared secret."],
            ["Rotation",
             "Manual — secrets expire and must be rotated.",
             "Automatic — short-lived tokens renewed by the platform."],
            ["Where it can run",
             "Anywhere (on-prem, container, laptop).",
             "Azure-hosted compute (Function App, Container App, AKS, App Service, VM)."],
            ["Token issuer",
             "v2 (login.microsoftonline.com/<tid>/v2.0) — set by the workload app's manifest.",
             "Default UAMI tokens are v1; production setups federate via FIC to get v2 (see §B.7)."],
            ["Best for",
             "Legacy, on-prem, local dev, drop-in replacement of static API keys.",
             "Greenfield Azure workloads; eliminates secret rotation entirely."],
        ],
        col_widths=[4, 6, 7],
    )

    add_callout(doc, "Confluent side is identical.",
                "Both patterns hit the same Confluent identity pool. The pool's filter is "
                "claims.tid == \"<tenant>\" && claims.aud in [<list-of-client-ids>]. To enable "
                "Pattern B for a workload that is already on Pattern A, append the UAMI's "
                "client ID to the pool's app_client_ids list — no new pool, no role-binding "
                "changes.")

    doc.add_page_break()

    # =========================================================================
    # Pattern A
    # =========================================================================
    add_heading(doc, "Pattern A — Client ID + Client Secret", 1)

    add_heading(doc, "A.1  When to use", 2)
    add_bullets(doc, [
        "Workload runs outside Azure (on-prem servers, third-party clouds, customer laptops).",
        "Drop-in replacement of an existing static Kafka API key with no infrastructure change.",
        "Local development where IMDS is not available.",
    ])

    add_heading(doc, "A.2  Flow", 2)
    add_numbered(doc, [
        "Workload reads tenant_id, client_id, and client_secret from configuration.",
        "MSAL calls Entra's /oauth2/v2.0/token with grant_type=client_credentials and "
        "scope=<client-id>/.default.",
        "Entra returns a v2 JWT with aud=<client-id>, tid=<tenant>, iss=https://"
        "login.microsoftonline.com/<tenant>/v2.0.",
        "Confluent.Kafka client passes the JWT as the SASL/OAUTHBEARER token, with the "
        "Confluent identityPoolId attached as an OAUTHBEARER extension.",
        "Confluent identity-pool filter matches claims.tid + claims.aud — authentication succeeds.",
    ])

    add_heading(doc, "A.3  Project layout", 2)
    add_code_block(doc,
        "DkDotnetCcClient/\n"
        "├── DkDotnetCcClient.csproj\n"
        "├── appsettings.json\n"
        "└── Program.cs")

    add_heading(doc, "A.4  DkDotnetCcClient.csproj", 2)
    add_code_block(doc, PATTERN_A_CSPROJ)

    add_heading(doc, "A.5  appsettings.json", 2)
    doc.add_paragraph(
        "Non-secret config only. The client_secret is supplied via the "
        "AZURE_CLIENT_SECRET environment variable so it never lands in source control."
    )
    add_code_block(doc, PATTERN_A_APPSETTINGS)

    add_heading(doc, "A.6  Program.cs", 2)
    add_code_block(doc, PATTERN_A_PROGRAM)

    add_heading(doc, "A.7  Run", 2)
    add_code_block(doc, PATTERN_A_RUN)

    add_heading(doc, "A.8  Pitfalls", 2)
    add_bullets(doc, [
        "Secret rotation. The client_secret expires (default 24 months); plan a rotation "
        "process and a delivery channel for the new value.",
        "Secret storage. Treat the secret like any other production credential — never "
        "commit, never log, never echo. Prefer Key Vault + a managed identity on the host "
        "if available; otherwise a sealed env var.",
        "Token caching. MSAL caches the access token in memory and reuses it until expiry. "
        "Do not fetch a new token per Kafka message — let the OAuthBearerTokenRefreshHandler "
        "fire only when librdkafka asks for one.",
    ])

    doc.add_page_break()

    # =========================================================================
    # Pattern B
    # =========================================================================
    add_heading(doc, "Pattern B — Function App + Managed Identity (Passwordless)", 1)

    add_heading(doc, "B.1  What “passwordless” actually means", 2)
    doc.add_paragraph(
        "A User-Assigned Managed Identity (UAMI) attached to a Function App is a real Entra "
        "service principal — it has its own object ID, its own client ID, its own tokens. "
        "What is different is how the Function App proves it owns that identity to Entra:"
    )
    add_bullets(doc, [
        "There is no shared secret. The Azure host runs an Instance Metadata Service "
        "(IMDS) endpoint at 169.254.169.254 that is reachable only from inside the "
        "Function App's sandbox.",
        "Code calls IMDS through the Azure.Identity SDK; the SDK exchanges that platform-"
        "assertion for a JWT signed by Entra.",
        "If a Function App's environment leaks (config dump, log scrape), nothing useful "
        "leaks — IMDS will not answer for anyone outside that host.",
        "Tokens are short-lived (~24 h) and auto-renewed by the platform. There is "
        "nothing to rotate, nothing to expire, nothing to deliver.",
    ])

    add_heading(doc, "B.2  When to use", 2)
    add_bullets(doc, [
        "Any workload hosted on Azure compute — Function App, Container App, App Service, "
        "AKS pod (with workload-identity), or a VM with a UAMI.",
        "Greenfield development where the operational cost of a secret-rotation process "
        "isn't justified.",
        "Replacing an existing Pattern A workload once it's been migrated to Azure.",
    ])

    add_heading(doc, "B.3  One-time setup (Azure / Entra admin)", 2)
    add_numbered(doc, [
        "Azure portal → Managed Identities → Create. "
        "Name: dk-confluent-{env}-{workload}-uami. Resource group: same as the Function App.",
        "Function App → Identity → User assigned → Add → select the UAMI.",
        "Open the UAMI's Overview blade and copy the Client ID (also called Application ID) — "
        "this is the GUID Confluent will trust.",
        "Append that Client ID to the workload's app_client_ids in "
        "terraform/live/<env>/workloads.json (the field is a list). "
        "Run terragrunt apply on that stack — the pool's filter updates in place.",
    ])

    add_callout(doc, "Why no Entra app registration?",
                "For Pattern B-Simple, the UAMI is the identity. You do not need to create "
                "a separate Entra app registration for the workload. The pool's filter "
                "trusts the UAMI's client ID directly. See §B.7 for the FIC variant where "
                "an app registration is reintroduced for stability across many UAMIs.")

    add_heading(doc, "B.4  Project layout", 2)
    add_code_block(doc,
        "DkFunctionAppMi/\n"
        "├── DkFunctionAppMi.csproj\n"
        "├── host.json\n"
        "├── local.settings.json    (gitignored — local dev only)\n"
        "├── Program.cs             (isolated worker bootstrap)\n"
        "└── KafkaProducerFunction.cs")

    add_heading(doc, "B.5  Source files", 2)

    add_code_block(doc, PATTERN_B_CSPROJ, caption="DkFunctionAppMi.csproj")
    add_code_block(doc, PATTERN_B_HOST, caption="host.json")
    add_code_block(doc, PATTERN_B_LOCAL_SETTINGS, caption="local.settings.json")
    add_code_block(doc, PATTERN_B_PROGRAM, caption="Program.cs")
    add_code_block(doc, PATTERN_B_FUNCTION, caption="KafkaProducerFunction.cs")

    add_heading(doc, "B.6  Deploy", 2)
    add_code_block(doc, PATTERN_B_DEPLOY)

    add_callout(doc, "v1 vs v2 token — verify before declaring victory.",
                "User-Assigned Managed Identities historically issue v1 access tokens "
                "(iss: https://sts.windows.net/<tid>/). The Confluent identity provider in "
                "this project is configured for v2 (iss: https://login.microsoftonline.com/"
                "<tid>/v2.0). Decode a sample token at jwt.ms — if iss is the v1 form, the "
                "Confluent broker will reject the JWT and you'll see SASL handshake failures "
                "with \"invalid_token\" / \"invalid issuer\" in the librdkafka logs. "
                "If you hit this, switch to the FIC variant in §B.7.",
                color=WARN, bg=WARN_BG)

    add_heading(doc, "B.7  Production variant — UAMI federated into a workload Entra app (FIC)", 2)
    doc.add_paragraph(
        "Use this variant when (a) §B.6 yields v1 tokens that Confluent rejects, or (b) you "
        "want one Entra app representing many Function/Container Apps so a single Confluent "
        "pool covers all of them."
    )

    add_heading(doc, "B.7.1  Setup", 3)
    add_numbered(doc, [
        "Keep the UAMI from §B.3 attached to the Function App.",
        "Create a workload Entra app registration: dk-confluent-{env}-{workload}-app. "
        "Set requestedAccessTokenVersion = 2 in the manifest (the project's PowerShell "
        "runbook does this automatically).",
        "On that app registration, add a Federated Identity Credential: "
        "issuer = https://login.microsoftonline.com/<tenant>/v2.0, "
        "subject = <UAMI's object ID>, "
        "audience = api://AzureADTokenExchange.",
        "Put the workload app's client ID (not the UAMI's) into app_client_ids "
        "for the Confluent pool. Apply.",
    ])

    add_heading(doc, "B.7.2  .NET — token acquisition", 3)
    doc.add_paragraph(
        "Replace the ManagedIdentityCredential constructor in KafkaProducerFunction with "
        "a ClientAssertionCredential that uses the UAMI as the assertion source:"
    )
    add_code_block(doc, PATTERN_B_FIC_CODE)

    doc.add_paragraph(
        "The resulting access token is issued by the workload app registration, so iss is "
        "guaranteed to be the v2 form and aud equals the workload app's client ID. The "
        "Function App still has no secret on disk — the UAMI is what closes the loop."
    )

    add_heading(doc, "B.8  Pitfalls", 2)
    add_bullets(doc, [
        "Local dev. ManagedIdentityCredential will fail on a developer laptop (no IMDS). "
        "For local testing, fall back to Pattern A or use DefaultAzureCredential plus an "
        "az-cli session with a user that the pool also trusts.",
        "Audience scoping. The scope passed to GetToken must be \"<client-id>/.default\" — "
        "this is how MSAL maps to a v2 audience equal to the client ID GUID. Using the "
        "api:// form will also work but produces an aud of api://<client-id>, which the "
        "Confluent pool filter does not accept by default in this project.",
        "Cold-start cost. The first token-acquisition call inside a Function invocation "
        "adds ~50–200 ms. Cache the credential in a static field; do not construct it "
        "per invocation.",
    ])

    doc.add_page_break()

    # =========================================================================
    # References
    # =========================================================================
    add_heading(doc, "Going further", 1)
    add_bullets(doc, [
        "Microsoft Entra workload identity — overview: "
        "https://learn.microsoft.com/azure/active-directory/workload-identities/workload-identities-overview",
        "Managed identities for Azure resources: "
        "https://learn.microsoft.com/entra/identity/managed-identities-azure-resources/overview",
        "Workload Identity Federation (FIC): "
        "https://learn.microsoft.com/entra/workload-id/workload-identity-federation",
        "Azure.Identity .NET reference: "
        "https://learn.microsoft.com/dotnet/api/overview/azure/identity-readme",
        "Confluent Cloud OAuth/OIDC for Kafka clients: "
        "https://docs.confluent.io/cloud/current/security/authenticate/workload-identities/identity-providers/oauth/clients/configure-kafka-client.html",
    ])

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
