"""
Generate docs/DK-Confluent-Entra-App-Runbook.docx — manual equivalent of the
New-WorkloadApps.ps1 PowerShell script, for DKP's Entra admin.

Usage:
    pip install python-docx --break-system-packages
    python docs/generate_entra_app_runbook.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

OUT = Path(__file__).parent / "DK-Confluent-Entra-App-Runbook.docx"

CONFLUENT_BLUE = RGBColor(0x00, 0x3E, 0x7E)
CODE_BG        = "F2F2F2"
MUTED          = RGBColor(0x55, 0x55, 0x55)
WARN           = RGBColor(0xB9, 0x60, 0x00)


def shade(element, fill: str) -> None:
    tc_pr = element._tc.get_or_add_tcPr() if hasattr(element, "_tc") \
            else element._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CONFLUENT_BLUE


def add_code_block(doc: Document, code: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    shade(cell, CODE_BG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.paragraphs[0].clear()
    for line in code.splitlines() or [""]:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Menlo"
        run.font.size = Pt(9)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text = "Field", "Value"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for i, (k, v) in enumerate(rows, start=1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(11)
    doc.add_paragraph()


def add_callout(doc: Document, label: str, text: str, color: RGBColor = CONFLUENT_BLUE) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, "EAF2FB")
    p = cell.paragraphs[0]
    run_label = p.add_run(f"{label}  ")
    run_label.font.bold = True
    run_label.font.color.rgb = color
    run_text = p.add_run(text)
    run_text.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DK Workload Identity")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = CONFLUENT_BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Entra ID App Registration Runbook")
    r.font.size = Pt(16); r.font.color.rgb = CONFLUENT_BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"v1.0 · {date.today():%B %Y} · Prepared by Confluent Professional Services")
    r.font.size = Pt(10); r.font.color.rgb = MUTED

    doc.add_paragraph()
    doc.add_paragraph()

    # -------- Purpose --------
    add_heading(doc, "Purpose", 1)
    doc.add_paragraph(
        "This runbook is the manual portal-based equivalent of the PowerShell script "
        "`New-WorkloadApps.ps1` shipped in the dk-workload-identity repository. It walks "
        "the DKP Entra-ID administrator through creating one app registration per "
        "(environment, domain, workload) so that each DKP workload can authenticate to "
        "Confluent Cloud with an Entra-issued JWT instead of a static Kafka API key."
    )
    doc.add_paragraph(
        "Each step below is repeated once per row in the accompanying spreadsheet "
        "`app-registrations.xlsx`. Field values come directly from that spreadsheet — "
        "the two documents are designed to be used side-by-side."
    )

    # -------- Who runs this --------
    add_heading(doc, "Who runs this", 1)
    doc.add_paragraph(
        "A DKP Entra-ID administrator (or anyone delegated the "
        "Application Administrator role in the DKP tenant) executes this runbook. "
        "The Confluent Professional Services team does not need access to the DKP tenant."
    )

    # -------- Prerequisites --------
    add_heading(doc, "Prerequisites", 1)
    doc.add_paragraph("Before you start, confirm you have:")
    for item in [
        "A browser logged in to the Azure portal as a user who holds the "
        "Application Administrator role (or equivalent) in the DKP Entra tenant "
        "(7bab0bc1-bb61-48d7-b2d4-79825c2ac6b8).",
        "The spreadsheet docs/app-registrations.xlsx open in a second window.",
        "A working copy of the dk-workload-identity repository cloned into DKP's Git.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # -------- What you'll create --------
    add_heading(doc, "What you will create", 1)
    doc.add_paragraph(
        "Fifteen Entra app registrations — one per row in the spreadsheet. "
        "Each app registration represents a single workload in a single environment "
        "and is paired 1:1 with a Confluent Cloud identity pool of the same name. "
        "The naming convention is:"
    )
    add_code_block(doc, "dk-confluent-{env}-{domain}-{workload}\n\nExample: dk-confluent-dev-mergerarb-madam")

    doc.add_paragraph(
        "Summary of what gets set per app, consistent across every row:"
    )
    add_kv_table(doc, [
        ("Sign-in audience",                       "Accounts in this organizational directory only (Single tenant)"),
        ("Redirect URI",                           "(leave blank — not needed for client credentials)"),
        ("Application ID URI",                     "api://<Application-(client)-ID>  (Azure sets this by default)"),
        ("Expose-an-API scope name & value",       "access_as_application"),
        ("Scope: admin consent display name",      "access_as_application"),
        ("Scope: admin consent description",       "Allow the application to access Kafka on behalf of the <workload> workload."),
        ("Scope: state",                           "Enabled"),
        ("Manifest: requestedAccessTokenVersion",  "2"),
    ])

    add_callout(doc, "Why v2 access tokens.",
                "Confluent's identity provider is configured with the v2 Entra issuer URL "
                "(https://login.microsoftonline.com/<tenant>/v2.0). v1 tokens use a different "
                "issuer (sts.windows.net) and will fail the Confluent trust check. Step 8 below "
                "flips the app to v2 via the manifest.")

    # -------- Procedure --------
    doc.add_page_break()
    add_heading(doc, "Procedure — repeat once per spreadsheet row", 1)
    add_callout(doc, "Tip.",
                "Do one full cycle (Steps 1–8) for a single row first, confirm it looks right, "
                "then proceed through the remaining rows.")

    steps = [
        ("Open the app-registrations screen",
         "Sign in to the Azure portal, open Microsoft Entra ID in the left rail, then "
         "select App registrations under Manage.",
         None),
        ("Start a new registration",
         "Click + New registration at the top of the list.",
         None),
        ("Fill in the registration form",
         "Use the values below (all from the spreadsheet row you are working on):",
         [
             ("Name",                    "Column E — App Display Name (e.g. dk-confluent-dev-mergerarb-madam)"),
             ("Supported account types", "Column F — Accounts in this organizational directory only (Single tenant)"),
             ("Redirect URI",            "Column G — leave blank"),
         ]),
        ("Register",
         "Click Register. The app's Overview page opens.",
         None),
        ("Record the IDs in the spreadsheet",
         "On the Overview page, copy the two GUIDs shown at the top into the highlighted "
         "columns of the spreadsheet:",
         [
             ("Application (client) ID", "Column O"),
             ("Object ID",                "Column P"),
         ]),
        ("Expose an API (Application ID URI)",
         "Left rail of the app → Expose an API → Set (next to 'Application ID URI'). "
         "Accept the default (api://<Application-(client)-ID>) and click Save. "
         "This value should match Column H template in the spreadsheet.",
         None),
        ("Add the access scope",
         "On the same Expose an API page, click + Add a scope and fill in:",
         [
             ("Scope name",                   "Column I — access_as_application"),
             ("Who can consent",              "Admins only"),
             ("Admin consent display name",   "Column J — access_as_application"),
             ("Admin consent description",    "Column K — Allow the application to access Kafka on behalf of the <workload> workload."),
             ("State",                        "Column L — Enabled"),
         ]),
        ("Flip the access-token version to 2 (manifest)",
         "Left rail of the app → Manifest. Find the property "
         "`requestedAccessTokenVersion` under the `api` object. Change it from "
         "`null` (or `1`) to `2`. Click Save at the top. This is Column M in the "
         "spreadsheet — it is always 2.",
         None),
    ]

    for i, (title, body, sub_rows) in enumerate(steps, start=1):
        add_heading(doc, f"Step {i}. {title}", 2)
        doc.add_paragraph(body)
        if sub_rows:
            add_kv_table(doc, sub_rows)

    # After all rows
    doc.add_page_break()
    add_heading(doc, "After the final row", 1)
    doc.add_paragraph(
        "Once every row in the spreadsheet has a filled-in Application (client) ID "
        "(Column O) and Object ID (Column P), the app-side setup is complete. Hand the "
        "spreadsheet back to the Confluent Professional Services team so they can:"
    )
    for item in [
        "Paste each Application (client) ID into terraform/live/<env>/workloads.json at the matching workload key.",
        "Run the terraform-workload GitHub Actions workflow for _org, dev, uat, and prd — in that order — to provision the Confluent identity pools + role bindings.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_callout(doc, "Service principal.",
                "Azure automatically provisions the service principal for each app as part of "
                "the portal-based New registration flow in Step 4 — no separate step is needed. "
                "(The PowerShell equivalent `New-MgApplication` does not do this, hence the "
                "extra `New-MgServicePrincipal` call in the script.)")

    # Troubleshooting
    add_heading(doc, "Troubleshooting", 1)
    add_kv_table(doc, [
        ("Confluent rejects the token with iss mismatch",
         "The app is still issuing v1 tokens. Re-check Step 8 — the manifest must have "
         "`\"requestedAccessTokenVersion\": 2` under the `api` object, saved."),
        ("Confluent rejects with aud mismatch",
         "Column O on the spreadsheet may have been mistyped into workloads.json. The pool "
         "filter expects the GUID alone (e.g. 72a90c20-…), not the api:// prefix."),
        ("'Insufficient privileges' when creating the app",
         "Your Entra account is not an Application Administrator in the DKP tenant. "
         "Ask the Global Administrator to grant the role or assign a delegated admin."),
        ("Scope name clashes with an existing one",
         "An earlier run of the runbook may have created the app already. Confirm Column O "
         "of the spreadsheet is still blank for this row; if not, this row has already been "
         "processed — skip it."),
    ])

    # Appendix
    doc.add_page_break()
    add_heading(doc, "Appendix A — Equivalent PowerShell (for reference only)", 1)
    doc.add_paragraph(
        "The repository under powershell/ ships an automated version of this same "
        "procedure. If your security posture permits running PowerShell against the DKP "
        "tenant with the Microsoft.Graph module, the commands below produce identical "
        "results without the manual portal steps."
    )
    add_code_block(doc,
        "# One-time connect\n"
        "Connect-MgGraph -TenantId 7bab0bc1-bb61-48d7-b2d4-79825c2ac6b8 `\n"
        "                -Scopes \"Application.ReadWrite.All\"\n"
        "\n"
        "# Create every app listed in workloads.<env>.json\n"
        "./powershell/scripts/New-WorkloadApps.ps1 `\n"
        "    -ConfigPath ./powershell/config/workloads.dev.json\n"
        "\n"
        "# Dump the app IDs into a JSON snippet for terragrunt\n"
        "./powershell/scripts/Export-WorkloadApps.ps1 `\n"
        "    -ConfigPath ./powershell/config/workloads.dev.json")

    return doc


if __name__ == "__main__":
    build().save(OUT)
    print(f"wrote {OUT}")
